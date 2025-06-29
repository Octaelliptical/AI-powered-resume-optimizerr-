import streamlit as st
import spacy
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import os

nlp = spacy.load("en_core_web_lg")
# Fix the path to be relative to the project root
skill_pattern_path = os.path.join("assets", "data", "jz_complete_patterns.jsonl")
nltk.download(['stopwords', 'wordnet'])

class ResumeAnalyzer:

    @staticmethod
    def remove_extra_spaces(text):
        words = text.split()
        cleaned_text = ' '.join(words)
        return cleaned_text

    @staticmethod
    def clean_text(text):
        review = re.sub(
            r'(@[A-Za-z0-9]+)|([^0-9A-Za-z \t])|(https?://\S+)|^rt',
            " ",
            text,
        )
        review = review.lower()
        review = review.split()
        lm = WordNetLemmatizer()
        review = [
            lm.lemmatize(word)
            for word in review
            if not word in set(stopwords.words("english"))
        ]
        review = " ".join(review)
        return review

    def extract_contact_number(text):
        pattern = r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"
        match = re.search(pattern, text)
        return match.group() if match else None

    def extract_email(text):
        pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
        match = re.search(pattern, text)
        return match.group() if match else None

    def extract_skills(text, nlp):
        doc = nlp(text)
        return [ent.text for ent in doc.ents if ent.label_ == "SKILL"]
    
    def unique_skills(text):
        return list(set(text))

    def extract_soft_skills(text,nlp):
        doc = nlp(text)
        return [ent.text for ent in doc.ents if ent.label_ == "SOFT_SKILL"]

    def extract_entities(text,nlp):
        entities = {label: [] for label in ["Job-Category", "ORG", "EDU", "GPE", "PK_ORG"]}
        doc = nlp(ResumeAnalyzer.clean_text(text))
        for ent in doc.ents:
            if ent.label_ in entities:
                entities[ent.label_].append(ent.text)
        return entities


    def extract_text_from_word_document(docx_file):
        try:
            doc = Document(docx_file)
            
            # Extract header text
            headers = [section.header for section in doc.sections]
            header_text = ""
            for header in headers:
                for paragraph in header.paragraphs:
                    header_text += paragraph.text + "\n"
            
            # Extract main content text
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            main_content_text = ' '.join(paragraphs)
            
            # Concatenate header and main content text
            full_text = header_text + "\n" + main_content_text
            
            return full_text
        except Exception as e:
            print(f"Error: {e}")
            return None

    def job_matching_algorithm(resume, job_description):
        if resume is None or job_description is None:
            return None
        text = [resume, job_description]
        cv = CountVectorizer()
        count_matrix = cv.fit_transform(text)
        try:
            similarity_scores = cosine_similarity(count_matrix)
            match_percentage = similarity_scores[0][1] * 100
            match_percentage = round(match_percentage, 2)
            return match_percentage
        except Exception as e:
            print(f"Error in job matching algorithm: {e}")
            return None

    @staticmethod
    def job_matching_feedback(match_percentage):
        if match_percentage >= 90:
            return "Master"
        elif 80 <= match_percentage < 90:
            return "Expert"
        elif 70 <= match_percentage < 80:
            return "Excellent"
        elif 50 <= match_percentage < 70:
            return "Good"
        elif 40 <= match_percentage < 50:
            return "Barely Passable"
        else:
            return "Needs Improvement"


    def display_demo_progress_bar():
        # Display the circular progress bar with a random number
        demo_percentage = 75
        demo_progress_bar_html = f"""  
        <div class="result-container">
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <div class="heading">
            </div>
            <div class="percentage-container">
                <div class="percentage-circle"></div>
                <span>{demo_percentage}%</span>
            </div>
            <h1>Demo</h1>
        </div>
        """
        # Define CSS to center align the content
        css = """
        <style>
        .result-container {
            margin: auto;

            width: 0%;
            text-align: center;
            justify-content: center;
        }
        </style>
        """
        st.markdown(css, unsafe_allow_html=True)
        st.markdown(demo_progress_bar_html, unsafe_allow_html=True)
